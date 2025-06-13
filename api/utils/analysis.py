import requests
import ipaddress
import shodan
from docx import Document
from datetime import datetime
import os

# Chaves de API
API_KEY_ABUSEIPDB = '39d2a0373bba7c859a8e0b23454efa5f393646946fce7c0ebdb4a52d51c9c1f40aa2bd499689bd4c'
API_KEY_VIRUSTOTAL = '5ad8844bec1ed8385dba214b9c9f1dbbebdc3a3bd4b67dd373f5613cb6d4c0ef'
API_KEY_SHODAN = 'Soa2pMLtSLgNU9X6BpphpeYhreRkNIMp'


def is_ip(address):
    try:
        ipaddress.ip_address(address)
        return True
    except ValueError:
        return False


def consultar_abuseipdb(ip):
    url = "https://api.abuseipdb.com/api/v2/check"
    headers = {'Key': API_KEY_ABUSEIPDB, 'Accept': 'application/json'}
    params = {'ipAddress': ip, 'maxAgeInDays': '90'}
    r = requests.get(url, headers=headers, params=params)
    if r.status_code == 200:
        return r.json()['data']
    return None


def consultar_virustotal(target, tipo):
    url = f"https://www.virustotal.com/api/v3/{tipo}/{target}"
    headers = {"x-apikey": API_KEY_VIRUSTOTAL}
    r = requests.get(url, headers=headers)
    if r.status_code == 200:
        data = r.json()
        attrs = data['data']['attributes']
        return {
            'reputation': attrs.get('reputation', 0),
            'last_analysis_stats': attrs.get('last_analysis_stats', {}),
            'last_analysis_date': attrs.get('last_analysis_date', 'N/A'),
        }
    return None


def consultar_shodan(ip):
    api = shodan.Shodan(API_KEY_SHODAN)
    try:
        res = api.host(ip)
        return {
            'org': res.get('org', 'N/A'),
            'os': res.get('os', 'N/A'),
            'ports': res.get('ports', []),
        }
    except:
        return None


def gerar_relatorio(target, is_ip_flag, dados_abuse, dados_vt, dados_shodan):
    doc = Document()
    doc.add_heading(f'Relatório - {target}', 0)

    if is_ip_flag and dados_abuse:
        doc.add_heading('AbuseIPDB', level=1)
        doc.add_paragraph(f"Abuse Score: {dados_abuse['abuseConfidenceScore']}")
        doc.add_paragraph(f"Total Reports: {dados_abuse['totalReports']}")

    doc.add_heading('VirusTotal', level=1)
    if dados_vt:
        stats = dados_vt['last_analysis_stats']
        doc.add_paragraph(f"Reputation: {dados_vt['reputation']}")
        doc.add_paragraph(f"Malicious: {stats.get('malicious',0)}")
        doc.add_paragraph(f"Suspicious: {stats.get('suspicious',0)}")
        doc.add_paragraph(f"Undetected: {stats.get('undetected',0)}")

    if is_ip_flag and dados_shodan:
        doc.add_heading('Shodan', level=1)
        doc.add_paragraph(f"Organization: {dados_shodan['org']}")
        doc.add_paragraph(f"OS: {dados_shodan['os']}")
        doc.add_paragraph(f"Open Ports: {dados_shodan['ports']}")

    nome_arquivo = os.path.join("reports", f"relatorio_{target.replace('/', '_')}.docx")
    doc.save(nome_arquivo)
    return nome_arquivo


def gerar_relatorio_hash(hash_str, dados_vt):
    doc = Document()
    doc.add_heading(f'Relatório - Análise de Hash', 0)

    doc.add_paragraph(f"Hash: {hash_str}")
    doc.add_paragraph(f"Data do Relatório: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph(" ")

    if dados_vt:
        stats = dados_vt.get("last_analysis_stats", {})
        doc.add_heading("VirusTotal", level=1)
        doc.add_paragraph(f"Reputação: {dados_vt.get('reputation', 0)}")
        doc.add_paragraph(f"Data da Última Análise: {dados_vt.get('last_analysis_date', 'N/A')}")
        doc.add_paragraph(f"Malicious: {stats.get('malicious', 0)}")
        doc.add_paragraph(f"Suspicious: {stats.get('suspicious', 0)}")
        doc.add_paragraph(f"Undetected: {stats.get('undetected', 0)}")
        doc.add_paragraph(f"Harmless: {stats.get('harmless', 0)}")

    nome_arquivo = os.path.join("reports", f"relatorio_hash_{hash_str[:10]}.docx")
    doc.save(nome_arquivo)
    return nome_arquivo


def enviar_arquivo_virustotal(file):
    url = "https://www.virustotal.com/api/v3/files"
    headers = {"x-apikey": API_KEY_VIRUSTOTAL}
    files = {"file": (file.filename, file.stream)}
    r = requests.post(url, headers=headers, files=files)
    if r.status_code == 200:
        return r.json()['data']['id']
    return None


def consultar_status_arquivo(scan_id):
    url = f"https://www.virustotal.com/api/v3/analyses/{scan_id}"
    headers = {"x-apikey": API_KEY_VIRUSTOTAL}
    r = requests.get(url, headers=headers)
    if r.status_code == 200:
        return r.json()['data']['attributes'].get('stats', {})
    return None
